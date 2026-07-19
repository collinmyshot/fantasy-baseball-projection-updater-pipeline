#!/usr/bin/env python3
from pathlib import Path
import csv
from datetime import datetime
from docx import Document
from docx.shared import Inches

ROOT = Path('data/processed/park_factors')
ARTICLE_DIR = ROOT / 'article'
FIG_DIR = ARTICLE_DIR / 'figs'
OUT_DOCX = ARTICLE_DIR / 'park_factor_article_2026_full.docx'

PARK_EVENTS_CSV = Path('data/manual/park_era_events.csv')
HOME_PARK_SOURCES_CSV = Path('data/manual/mlb_home_parks_2026_verified_sources.csv')
DEF_SOURCES_CSV = Path('data/manual/team_defense_2015_2025_sources.csv')


def read_csv(path):
    with open(path, newline='', encoding='utf-8') as f:
        return list(csv.DictReader(f))


def to_float(x, default=0.0):
    try:
        return float(x)
    except Exception:
        return default


def to_int(x, default=0):
    try:
        return int(round(float(x)))
    except Exception:
        return default


def weighted_mean(rows, val_key, w_key):
    num = 0.0
    den = 0.0
    for r in rows:
        v = to_float(r.get(val_key, 'nan'), default=float('nan'))
        w = to_float(r.get(w_key, 'nan'), default=float('nan'))
        if v == v and w == w:
            num += v * w
            den += w
    return num / den if den else float('nan')


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)


def find_meta(meta_rows, key, default=''):
    for r in meta_rows:
        if r.get('key') == key:
            return r.get('value', default)
    return default


def top_n(rows, key, n=10):
    return sorted(rows, key=lambda r: to_float(r.get(key, 0)), reverse=True)[:n]


def bot_n(rows, key, n=10):
    return sorted(rows, key=lambda r: to_float(r.get(key, 0)))[:n]


def team_set(rows):
    return {r['Team'] for r in rows}


def insert_figure(doc, filename, title=None):
    fp = FIG_DIR / filename
    if fp.exists():
        if title:
            doc.add_paragraph(title)
        doc.add_picture(str(fp), width=Inches(6.8))


def main():
    overall = read_csv(ROOT / 'park_factors_savant_style_clean_2026.csv')
    val_sum = read_csv(ROOT / 'validation_summary.csv')
    invariance = read_csv(ROOT / 'invariance_checks.csv')
    weights = read_csv(ROOT / 'park_factors_savant_style_weights.csv')
    known = read_csv(ROOT / 'park_factor_known_effects_2026.csv')
    meta = read_csv(ROOT / 'run_metadata.csv')
    ext_sum = read_csv(ROOT / 'external_validation_summary.csv')
    ext_detail = read_csv(ROOT / 'external_validation_detail.csv')

    park_events = read_csv(PARK_EVENTS_CSV) if PARK_EVENTS_CSV.exists() else []
    home_park_sources = read_csv(HOME_PARK_SOURCES_CSV) if HOME_PARK_SOURCES_CSV.exists() else []
    defense_sources = read_csv(DEF_SOURCES_CSV) if DEF_SOURCES_CSV.exists() else []

    mean_rmse_model = weighted_mean(val_sum, 'rmse_model', 'n_park_half')
    mean_rmse_zero = weighted_mean(val_sum, 'rmse_zero', 'n_park_half')
    mean_rmse_prev = weighted_mean(val_sum, 'rmse_prev', 'n_park_half')
    mean_corr = weighted_mean(val_sum, 'corr_model_vs_realized', 'n_park_half')
    mean_corr_prev = weighted_mean(val_sum, 'corr_prev_vs_realized', 'n_park_half')
    mean_slope = weighted_mean(val_sum, 'calibration_slope', 'n_park_half')

    inv = {r['metric']: to_float(r.get('value', 'nan'), default=float('nan')) for r in invariance}

    rows_modeled = find_meta(meta, 'rows_modeled', 'NA')
    seasons_modeled = find_meta(meta, 'seasons_modeled', '')
    excluded = find_meta(meta, 'exclude_seasons', '')

    overall_sorted = sorted(overall, key=lambda r: to_float(r['Overall Park Factor']), reverse=True)
    top10 = overall_sorted[:10]
    bottom10 = list(reversed(overall_sorted[-10:]))

    hb = [r for r in known if r.get('Analysis') == 'HR vs BACON Gap']
    hb_sorted = sorted(hb, key=lambda r: abs(to_float(r.get('Difference', '0'))), reverse=True)

    split = [r for r in known if r.get('Analysis') == '1H vs 2H Overall PF Gap']
    split_sorted = sorted(split, key=lambda r: abs(to_float(r.get('Difference', '0'))), reverse=True)

    # External validation metrics map.
    ext = {r['metric']: r for r in ext_sum}

    # Top/bottom overlap checks for context in validation section.
    usable = [r for r in ext_detail if r.get('savant_index_woba', '') not in ('', None)]
    our_top10 = team_set(top_n(usable, 'our_overall_pf', 10))
    our_bot10 = team_set(bot_n(usable, 'our_overall_pf', 10))
    sv_top10 = team_set(top_n(usable, 'savant_index_woba', 10))
    sv_bot10 = team_set(bot_n(usable, 'savant_index_woba', 10))

    fg_top10 = team_set(top_n(ext_detail, 'fg_basic5yr_pf', 10))
    fg_bot10 = team_set(bot_n(ext_detail, 'fg_basic5yr_pf', 10))
    our_top10_all = team_set(top_n(ext_detail, 'our_overall_pf', 10))
    our_bot10_all = team_set(bot_n(ext_detail, 'our_overall_pf', 10))

    savant_top_overlap = len(our_top10 & sv_top10)
    savant_bot_overlap = len(our_bot10 & sv_bot10)
    fg_top_overlap = len(our_top10_all & fg_top10)
    fg_bot_overlap = len(our_bot10_all & fg_bot10)

    # Park-era summaries.
    park_event_rows = []
    for r in park_events:
        park_event_rows.append([
            r.get('team', ''),
            r.get('event_type', ''),
            r.get('era_suffix', ''),
            r.get('start_date', ''),
            r.get('end_date', '') if r.get('end_date') else 'active',
            r.get('source_primary', '')
        ])

    doc = Document()
    doc.add_heading('Building Fantasy-Forward Park Factors: A Residual-Based Park Model for Streaming Decisions', level=0)
    p = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M %Z")}')
    p.runs[0].italic = True

    doc.add_heading('Ugh, Another Park Factor Model?', level=1)
    doc.add_paragraph(
        'Park factors (PFs) are something we all think about when evaluating hitters and pitchers. '
        'Colorado and Cincinnati are feared starts for pitchers, and for good reason. '
        'But once you move past the obvious spots, PF math can feel opaque and sometimes disconnected from what we, as fantasy players, actually experience week to week. '
        'That is exactly why I built this model.'
    )
    doc.add_paragraph(
        'Baseball Savant sets the standard for public park factors '
        '(https://baseballsavant.mlb.com/leaderboard/statcast-park-factors), but the table includes plenty of categories that are either less relevant for most fantasy formats '
        '(2B, 3B, xBACON) or less intuitive as park-driven levers in day-to-day fantasy decisions (OBP, Hard Hit). '
        'FanGraphs park factors (https://www.fangraphs.com/tools/guts?type=pf) are also useful, but are less commonly referenced in fantasy circles and are built with a different objective. '
        'Both models are trying to capture real-life value. Fantasy managers, especially streamers, usually need something narrower: '
        'how much a park changes BABIP-type hit risk (WHIP and innings length) and how much it changes home-run blow-up risk.'
    )

    doc.add_heading('So How Does This Model Work?', level=1)
    doc.add_paragraph('My methodology has three main pillars:')
    doc.add_paragraph(
        'Actual minus expected contact outcomes: if a park repeatedly over- or under-performs expected contact quality over a long sample, '
        'the park itself is a likely driver.',
        style='List Number'
    )
    doc.add_paragraph(
        'Weather and baseball physics: this model controls for game-level weather when available (temperature, wind, humidity) and year-level drag effects, '
        'so we do not blame the park for what is really atmospheric or ball-level noise.',
        style='List Number'
    )
    doc.add_paragraph(
        'Team defense: strong defenses can suppress observed BABIP independent of park geometry, so defense has to be controlled or park estimates get biased.',
        style='List Number'
    )
    doc.add_paragraph(
        'Concretely, this model incorporates every batted-ball event (BBE) from the start of the Statcast era in 2015 through the end of the 2025 regular season '
        '(2020 excluded because it was such a confounding year, both in real life and from a data perspective; source endpoint: https://baseballsavant.mlb.com/statcast_search/csv). '
        f'That leaves {to_int(rows_modeled):,} BBEs across {seasons_modeled}. '
        'The primary residual is wOBAcon minus xwOBAcon, with component residuals for BACON and HR (XBH is retained in the display blend). '
        'Defense is included at the team-season level using an even blend of OAA, UZR, and DRS '
        '(OAA source: https://baseballsavant.mlb.com/leaderboard/outs_above_average; DRS/UZR source: https://www.fangraphs.com/leaders/major-league?stats=fld). '
        'I also split environmental effects into 1H (March-June) and 2H (July-September) to capture broad seasonal shifts.'
    )
    doc.add_paragraph(
        'Nerd stuff alert: the model is hierarchical, with random effects for park-era, park-era-half, batter-season, pitcher-season, batting-team-season, and fielding-team-season. '
        'That is the core move that separates park context from roster quality. Invariance checks are reassuringly small: correlation between estimated park effect and home-team offense is '
        f'{inv.get("corr_park_effect_vs_home_team_xwoba_con", float("nan")):.3f}, and correlation with home-team defense is '
        f'{inv.get("corr_park_effect_vs_home_team_defense", float("nan")):.3f}. '
        'TL;DR: once talent and defense are controlled, repeatable venue-era residuals are treated as park signal.'
    )
    doc.add_paragraph(
        'One structural choice worth calling out: this build explicitly splits park eras around known moves, temporary homes, and meaningful dimension changes '
        '(data/manual/park_era_events.csv; full event log in Appendix A). Public models still account for these effects in aggregate, but they usually do not expose boundary-aware era splits directly '
        'in the output table. In practice, that means rolling windows like Savant\'s 3-year view can blend pre-change and post-change environments in a single park line.'
    )
    doc.add_paragraph(
        'For 2026 outputs, we also maintain a verified home-park map for all 30 teams in data/manual/mlb_home_parks_2026_verified_sources.csv, '
        'using MLB team pages (https://www.mlb.com/team) with secondary checks against current stadium listings.'
    )

    doc.add_heading('How the Scores Are Built (and What RMSE Means)', level=1)
    doc.add_paragraph(
        'RMSE is the quick scorecard for "how wrong are we, on average, in holdout data?" In this case it is the square root of weighted mean squared prediction error for park-half residuals in rolling validation. '
        'We train on trailing windows and predict the next season to test forward stability, not just in-sample fit. '
        f'Weighted holdout RMSE values: model={mean_rmse_model:.4f}, zero-baseline={mean_rmse_zero:.4f}, prior-park-baseline={mean_rmse_prev:.4f}. '
        f'Weighted correlation between predicted and realized holdout effects is {mean_corr:.3f} (prior-park baseline {mean_corr_prev:.3f}), with calibration slope {mean_slope:.3f}. '
        'For display, we normalize everything to a 100-centered index where 100 is league-average context. '
        'The final Overall PF then combines BACON, HR, and XBH with your chosen fantasy blend (0.45 / 0.35 / 0.20), so the metric reflects both hit volume and damage quality.'
    )

    add_table(
        doc,
        ['Component', 'Weight'],
        [[w['component'], f"{to_float(w['weight']):.2f}"] for w in weights]
    )

    insert_figure(doc, 'validation_rmse.png', 'Rolling holdout RMSE by season')
    insert_figure(doc, 'pred_vs_realized.png', 'Predicted versus realized holdout park-half residuals')
    insert_figure(doc, 'overall_pf_bar.png', 'Overall PF index (100 = league average)')

    doc.add_heading('2026 Team Results Snapshot', level=2)
    add_table(
        doc,
        ['Rank', 'Team', 'Park', 'Years', 'Overall PF', 'BACON PF', 'HR PF', 'Total BBE'],
        [[
            r.get('Rank', ''), r.get('Team', ''), r.get('Park', ''), r.get('Years', ''),
            f"{to_float(r.get('Overall Park Factor')):.2f}",
            f"{to_float(r.get('BACON Park Factor')):.2f}",
            f"{to_float(r.get('HR Park Factor')):.2f}",
            f"{to_int(r.get('Total BBE'))}"
        ] for r in top10]
    )

    add_table(
        doc,
        ['Rank', 'Team', 'Park', 'Years', 'Overall PF', 'BACON PF', 'HR PF', 'Total BBE'],
        [[
            r.get('Rank', ''), r.get('Team', ''), r.get('Park', ''), r.get('Years', ''),
            f"{to_float(r.get('Overall Park Factor')):.2f}",
            f"{to_float(r.get('BACON Park Factor')):.2f}",
            f"{to_float(r.get('HR Park Factor')):.2f}",
            f"{to_int(r.get('Total BBE'))}"
        ] for r in bottom10]
    )

    doc.add_heading('How It Compares to Savant and FanGraphs', level=1)
    doc.add_paragraph(
        'To sanity-check this against models fantasy players already trust, we validated directly against Savant and FanGraphs park outputs using matched team-level indices. '
        'For Savant (https://baseballsavant.mlb.com/leaderboard/statcast-park-factors), our Overall PF versus Savant index_wOBA yields '
        f'r={to_float(ext["savant_overall_vs_index_woba"]["pearson_r"]):.3f} across n={to_int(ext["savant_overall_vs_index_woba"]["n_parks"])} parks; '
        f'BACON alignment is r={to_float(ext["savant_bacon_vs_index_bacon"]["pearson_r"]):.3f}; HR alignment is r={to_float(ext["savant_hr_vs_index_hr"]["pearson_r"]):.3f}. '
        'For FanGraphs Guts park factors (https://www.fangraphs.com/tools/guts?type=pf), using Basic (5yr) as the primary comparator as requested, '
        f'Overall PF alignment is r={to_float(ext["fg_overall_vs_basic5yr"]["pearson_r"]):.3f} across n={to_int(ext["fg_overall_vs_basic5yr"]["n_parks"])} parks; '
        f'HR-to-HR alignment is r={to_float(ext["fg_hr_vs_hr"]["pearson_r"]):.3f}. '
        'By rank overlap, our top-10 and bottom-10 overlap with Savant is '
        f'{savant_top_overlap}/10 and {savant_bot_overlap}/10, and with FG Basic (5yr) is {fg_top_overlap}/10 and {fg_bot_overlap}/10. '
        'Bottom line: strong directional agreement with both public models, but with a more fantasy-usable component split by design.'
    )

    add_table(
        doc,
        ['Validation Metric', 'Parks', 'Pearson r', 'Spearman rho'],
        [[
            r['metric'],
            r['n_parks'],
            f"{to_float(r['pearson_r']):.3f}",
            f"{to_float(r['spearman_rho']):.3f}"
        ] for r in ext_sum]
    )

    doc.add_heading('Conclusions: What This Means for Streaming', level=1)
    doc.add_paragraph(
        'The core takeaway is that "offense-friendly park" is not one thing. Some parks boost hit conversion more than long-ball damage, while others do the opposite. '
        'That matters directly for streaming because pitcher archetypes are not interchangeable: fly-ball arms are more exposed in HR-tilted parks, while contact managers are more exposed in BACON-tilted parks.'
    )
    doc.add_paragraph(
        'A second takeaway is timing. The 1H/2H split surfaces repeatable seasonal drift that a single full-season park number can hide, which is useful when borderline streaming calls depend on thin margins. '
        'Operationally, this model is best used as a context layer alongside skill and opponent quality: it should sharpen close decisions, not override large talent gaps.'
    )
    doc.add_paragraph(
        'Finally, validation suggests this framework is directionally aligned with Savant and FanGraphs while adding fantasy-specific component clarity. '
        'That is the point of the build: keep the signal those models capture, then make the output easier to act on for weekly roster and streaming decisions.'
    )

    insert_figure(doc, 'hr_minus_bacon.png', 'HR PF minus BACON PF (positive = HR-friendlier than hit-friendlier)')
    insert_figure(doc, 'half_split_movers.png', 'Largest 1H versus 2H movers')

    add_table(
        doc,
        ['Analysis', 'Team', 'Park', 'Difference', 'Abs Difference', 'Years', 'Total BBE'],
        [[
            r.get('Analysis', ''),
            r.get('Team', ''),
            r.get('Park', ''),
            f"{to_float(r.get('Difference')):.2f}",
            f"{to_float(r.get('Abs Difference')):.2f}",
            r.get('Years', ''),
            f"{to_int(r.get('Total BBE'))}"
        ] for r in (hb_sorted[:12] + split_sorted[:12])]
    )

    doc.add_heading('What This Model Is (and Is Not)', level=1)
    doc.add_paragraph(
        f'Assumptions and exclusions: 2020 is excluded ({excluded}); park-era tagging quality is bounded by public reporting and source logs; '
        'weather/drag covariates are included but noisy; 1H/2H is intentionally coarse; and spray-angle interaction terms were not promoted to the production fixed-effect set for stability reasons. '
        'The goal is simple and practical: estimate true park run environment independent of team quality, then expose component-specific risk (especially BACON and HR) to make streaming calls more confident.'
    )

    doc.add_heading('Appendix A: Park-Era Event Log', level=1)
    doc.add_paragraph(
        'Full park-era event definitions used to split venue contexts, including move dates, wall-change windows, and primary source URLs.'
    )
    if park_event_rows:
        add_table(
            doc,
            ['Team', 'Event Type', 'Era Suffix', 'Start', 'End', 'Primary Source URL'],
            park_event_rows
        )

    doc.add_heading('Appendix B: Equations and Model Formulae', level=1)
    doc.add_paragraph('Event-level residual definitions:')
    doc.add_paragraph('resid_i = wOBAcon_i - xwOBAcon_i', style='List Bullet')
    doc.add_paragraph('bacon_resid_i = hit_on_contact_i - xBA_on_contact_i', style='List Bullet')
    doc.add_paragraph('hr_resid_i = HR_on_contact_i - xHR_on_contact_i', style='List Bullet')
    doc.add_paragraph('xbh_resid_i = XBH_on_contact_i - xXBH_on_contact_i', style='List Bullet')

    doc.add_paragraph('Primary mixed-effects model (conceptual fixed+random form):')
    doc.add_paragraph(
        'resid ~ half + temp + wind_speed + humidity + drag + defense_composite '
        '+ (1 | park_era_id) + (1 | park_era_half_id) + (1 | batter_season_id) '
        '+ (1 | pitcher_season_id) + (1 | fielding_team_season_id) + (1 | batting_team_season_id)',
        style='List Bullet'
    )

    doc.add_paragraph('Component models (same random structure, outcome swapped):')
    doc.add_paragraph('bacon_resid ~ same fixed/random structure', style='List Bullet')
    doc.add_paragraph('hr_resid ~ same fixed/random structure', style='List Bullet')
    doc.add_paragraph('xbh_resid ~ same fixed/random structure', style='List Bullet')

    doc.add_paragraph('Index scaling and aggregation:')
    doc.add_paragraph('PF_component_idx = 100 * (1 + delta_component / baseline_component)', style='List Bullet')
    doc.add_paragraph(
        'Overall_PF_idx = 0.45 * BACON_idx + 0.35 * HR_idx + 0.20 * XBH_idx',
        style='List Bullet'
    )
    doc.add_paragraph('RMSE = sqrt(weighted_mean((realized - predicted)^2, n_bbe))', style='List Bullet')

    doc.add_heading('Inline Sources', level=1)
    src_lines = [
        'Statcast search endpoint: https://baseballsavant.mlb.com/statcast_search/csv',
        'Savant park factors page: https://baseballsavant.mlb.com/leaderboard/statcast-park-factors',
        'Savant OAA leaderboard: https://baseballsavant.mlb.com/leaderboard/outs_above_average',
        'FanGraphs Guts park factors: https://www.fangraphs.com/tools/guts?type=pf',
        'FanGraphs fielding leaders (DRS/UZR): https://www.fangraphs.com/leaders/major-league?stats=fld',
        'MLB team pages (home park verification): https://www.mlb.com/team',
        f'Park era source rows included: {len(park_events)} (data/manual/park_era_events.csv)',
        f'2026 home park verification rows included: {len(home_park_sources)} (data/manual/mlb_home_parks_2026_verified_sources.csv)',
        f'Defense source rows included: {len(defense_sources)} (data/manual/team_defense_2015_2025_sources.csv)'
    ]
    for s in src_lines:
        doc.add_paragraph(s, style='List Bullet')

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f'Wrote {OUT_DOCX}')


if __name__ == '__main__':
    main()
