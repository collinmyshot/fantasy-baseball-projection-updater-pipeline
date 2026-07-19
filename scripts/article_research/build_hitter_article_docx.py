#!/usr/bin/env python3
"""Build an editable Word version of the Hitter Streamonator methodology doc."""
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1f, 0x35, 0x56)
GOLD = RGBColor(0xb0, 0x86, 0x37)
GRAY = RGBColor(0x77, 0x77, 0x77)
DARK = RGBColor(0x22, 0x22, 0x22)
CAVEAT = RGBColor(0x4a, 0x2f, 0x10)

doc = Document()

# ---- base style -----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(10)
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
    setattr(section, m, Inches(1))


def shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_border_bottom(cell, color="D5CCB4", sz="12"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    tcPr.append(borders)


def add_hyperlink(paragraph, url, text, color="1F3556", underline=True, size_pt=None):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if color:
        c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    if underline:
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    if size_pt:
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size_pt * 2))); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def para(segments, space_after=10, space_before=0, align=None, indent=None, style=None):
    """segments: list of (text, bold, italic, color, size_pt) tuples, or a plain string."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if isinstance(segments, str):
        segments = [(segments, False, False, None, None)]
    for seg in segments:
        text, bold, italic, color, size = (list(seg) + [None] * 5)[:5]
        r = p.add_run(text)
        r.bold = bool(bold)
        r.italic = bool(italic)
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
        if size:
            r.font.size = Pt(size)
    return p


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    return p


def byline(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GRAY
    return p


def h2(part_label, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(part_label.upper())
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = GOLD
    p2 = doc.add_heading(level=1)
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(10)
    r2 = p2.add_run(title)
    r2.bold = True; r2.font.size = Pt(17); r2.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    r2.font.name = "Calibri"
    pPr = p2._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12"); bottom.set(qn("w:color"), "E8E4D8")
    pbdr.append(bottom); pPr.append(pbdr)
    return p2


def h3(text):
    p = doc.add_heading(level=2)
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True; r.italic = True; r.font.size = Pt(12.5); r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    r.font.name = "Calibri"
    return p


def pull_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "24"); left.set(qn("w:color"), "C8A96E")
    left.set(qn("w:space"), "8")
    pbdr.append(left); pPr.append(pbdr)
    r = p.add_run(text)
    r.bold = True; r.italic = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
    return p


def caption(text):
    return para([(text, False, True, "777777", 9)], space_after=16)


def table(headers, rows, widths_in, header_bold_first_col=True, note_span=None):
    """rows: list of lists of strings (or (text, colspan) for spans). widths_in: list of inches."""
    ncols = len(headers)
    t = doc.add_table(rows=1, cols=ncols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(htext)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        shade_cell(hdr[i], "F2EDE2")
        hdr[i].width = Inches(widths_in[i])
    for row in rows:
        cells = t.add_row().cells
        ci = 0
        for j, val in enumerate(row):
            if ci >= ncols:
                break
            span = 1
            if isinstance(val, tuple):
                val, span = val
            cell = cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
            if j == 0:
                r.bold = True
            if span > 1:
                merged = cell
                for k in range(1, span):
                    merged = merged.merge(cells[ci + k])
            cell.width = Inches(widths_in[ci])
            ci += span
    for row in t.rows:
        for cell in row.cells:
            cell.width = Inches(widths_in[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def caveat_item(lead, rest):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    r1 = p.add_run(lead + " ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = CAVEAT
    r2 = p.add_run(rest)
    r2.font.size = Pt(10); r2.font.color.rgb = CAVEAT
    return p


BOLD, ITAL, KEY = "bold", "ital", "key"  # readability markers, unused directly

# ============================================================================
# TITLE
# ============================================================================
h1("Hitter Streamonator: Methodology")
byline("Methodology Series  ·  How (and whether) hitters can be streamed")

para([
    ("The SP Streamonator collapses a start into one composite score because a start is one big, "
     "self-contained event. Hitters do not cooperate: a single game is three to five plate appearances "
     "of near-pure noise, and value is smeared across five roto categories that move independently. "
     "The honest headline has two halves. On any single decision (pick the better of two similarly "
     "projected hitters) these models are barely better than a coin flip, because a hitter’s week is "
     "roughly 23 plate appearances of mostly noise. But that thin per-decision edge is real, and it "
     "compounds: summed across a lineup’s worth of picks, the models’ top five out-produce the "
     "preseason projection’s top five in 59 to 66 percent of weeks, category by category, by "
     "measurable margins of production. This is a modest, conditional tool, strongest for home runs and "
     "weakest for waiver-tier batting average, and every figure below is stated that way. Validation is "
     "out of sample throughout (leave-one-season-out, 2016–2025, roughly 368,000 hitter-games); where "
     "a number fails to clear its confidence interval, that is said plainly.", False, False, None, None),
])

table(
    ["Category (weekly)", "Single pick vs a coin flip", "Top-5 beats projection", "On the streamable pool"],
    [
        ["Batting average", "50–52%", "66% of weeks [61, 72]", "53% [47, 60]"],
        ["Stolen bases", "51–53%", "62% [57, 67]", "65% [59, 70]"],
        ["Home runs", "50–53%", "61% [56, 66]", "63% [58, 68]"],
        ["Overall (5×5 blend)", "—", "61% [55, 67]", "not tested"],
    ],
    [1.9, 1.7, 1.7, 1.4],
)
caption("Read left to right, the metric gets more favorable and more honest at once. A single pick "
        "between two similar hitters is close to a coin flip; a top-five bundle of those picks tips the "
        "week most of the time; the streamable pool (players outside the top 150 draft picks, or "
        "undrafted) is where the tool actually lives. Note the honest exception up front: the "
        "batting-average edge concentrates among rostered players and does not clear significance on "
        "the waiver tier.")

# ============================================================================
# PART I
# ============================================================================
h2("Part I", "The Noise Floor, and Why Steamer Is the Bar")

para("Before building anything, the ceiling was measured. Using a properly scaled preseason projection "
     "(Steamer, converted to a weekly 5×5 value and compared against what hitters actually did), how "
     "much of a hitter’s realized weekly value can a projection explain?")

pull_quote("Projections explain about 8.7 percent of the week-to-week variance in hitter value. The "
           "remaining 91 percent behaves like a coin flipped 23 times.")

para([
    ("An aggregate number understates how hard the real decision is, so consider the question a "
     "streamer actually faces: two hitters with similar projections, pick the one who has the better "
     "week. Conditioning on a realistic similarity (a gap of ten ranks in the projected pool), the "
     "projection identifies the better week in ", False, False, None, None),
    ("50 to 53 percent of pairs", True, False, "111111", None),
    (" — a coin flip with a faint thumb on the scale. An all-pairs figure looks better, selecting "
     "the better player-week about 58 percent of the time, but that number is inflated by easy "
     "comparisons between stars and scrubs; nobody is sitting Shohei Ohtani for a waiver-wire pickup.",
     False, False, None, None),
])

para([
    ("This is not a failure of projections; it is the physics of a 23-plate-appearance sample. It is "
     "also precisely why the tool can exist: ", False, False, None, None),
    ("if projections cleanly separated similar hitters, there would be nothing left to add.",
     True, False, "111111", None),
    (" The entire enterprise is finding the sliver of signal a season-long projection cannot see (the "
     "park tonight, the pitcher on the mound, the batting-order slot this week) and proving that sliver "
     "survives the noise.", False, False, None, None),
])

para("Why is the preseason projection the baseline to beat, rather than something fancier? Because the "
     "simple alternatives are worse. Ranking hitters by their last 14 or last 30 days, or picking "
     "between comparable options at random, loses to the projection consistently; at the weekly horizon "
     "the engine beats random selection in 83 to 98 percent of weeks and beats recency ranking in 80 to "
     "91 percent, far more often than it beats the projection itself. The projection is the strongest "
     "simple strategy available, so the question this tool answers is: given that Steamer is the best "
     "free lunch, what beats Steamer?")

para("Two housekeeping notes before the models. First, the target: each hitter-week is scored on actual "
     "Mon–Sun production, z-scored within the week across the startable pool on the standard roto "
     "5×5 categories (AVG, R, RBI, HR, SB), with average valued as hits above what a league-average "
     "hitter would post in the same at-bats; the z-scoring approach is described on the Hitter Valuation "
     "page. Second, the data: the engine is built on roughly 368,000 started hitter-games from "
     "2016–2025 (2020 excluded), pulled from the MLB Stats API. The atomic unit is one hitter, one "
     "game, with its park, opposing starter, that starter’s handedness, the defending team, and the "
     "hitter’s batting-order slot; a series or a full week is a sum of these atoms. The spine "
     "reconciles against the official league record exactly (2024 totals: 5,453 HR and 3,617 SB on both "
     "sides, to the digit).")

# ============================================================================
# PART II
# ============================================================================
h2("Part II", "Clearing the Deck: Two Popular Instincts, Tested")

h3("The hot bat, split into what matters and what does not")

para("Trailing performance was tested at every window (last 3, 7, 14, 30, and 60 days) as an addition "
     "to the projection. The verdict splits cleanly by category:")

para([
    ("Where recent form matters: ", True, False, "111111", None),
    ("power, and to a lesser degree steals. An in-season home-run surge carries real signal; the "
     "heavily shrunk to-date HR rate earns a ×1.32 swing in the HR model, rivaling the projection "
     "itself. Power changes mid-season are often real changes (health, bat speed, swing adjustments), "
     "and the model listens. The to-date steal rate helps too, and is the single largest lever in the "
     "SB model (×1.65), because a hitter who has started running is a hitter whose team has given "
     "him the green light.", False, False, None, None),
])

para([
    ("Where it does not: ", True, False, "111111", None),
    ("batting average, and above all the decision itself. A hot average adds essentially nothing that "
     "survives shrinkage (its to-date term is ×1.05, a whisper). And ranking hitters purely by "
     "recent form is the worst strategy tested in this entire project: worse than the projection, worse "
     "than the engine, beaten by the engine in 80 to 91 percent of weeks. “Start whoever is "
     "hot” is the instinct the data most firmly rejects.", False, False, None, None),
])

h3("Platoon splits: real for the league, unknowable for the player")

para("The platoon matchup feels decisive, and at the league level it is real. Here is the "
     "league-average grid, 2015–2025, from FanGraphs splits (switch hitters counted on the side "
     "they actually bat in each matchup):")

table(["AVG", "vs LHP", "vs RHP"],
      [["LHB", ".236", ".250"], ["RHB", ".256", ".246"]],
      [1.5, 1.5, 1.5])

table(["HR per 600 PA", "vs LHP", "vs RHP"],
      [["LHB", "14.3", "19.2"], ["RHB", "19.7", "18.6"]],
      [1.5, 1.5, 1.5])
caption("2.1 million plate appearances, 2015–2025. The average platoon effect is asymmetric: a "
        "lefty bat with the platoon advantage hits home runs at a much higher rate (19.2 versus 14.3 "
        "per 600), while a righty bat gains far less (19.7 versus 18.6). This raw split looks enormous, "
        "but it conflates the matchup with everything correlated to it.")

para([
    ("So platoon advantages exist. The problem is twofold. First, attributing them to individuals: a "
     "hitter’s true personal split is one of the slowest-stabilizing quantities in baseball. "
     "Research at The Hardball Times ", False, False, None, None),
    ("puts the requirement at roughly 2,000 plate appearances against left-handed pitching, which, at "
     "the league’s ~28 percent lefty exposure, means 6,000 to 7,000 total PA: more than ten seasons "
     "of everyday play. For practically every hitter in a streaming pool, the individual split never "
     "becomes separable from the league-wide one. Second, and more sobering: once the projection is in "
     "the model (it already embeds each hitter’s typical handedness exposure), the marginal effect "
     "of tonight’s platoon advantage on home runs is only ×1.05. The raw grid is dominated by "
     "which hitters get the platoon advantage, not by the advantage itself. The engine therefore gives "
     "every hitter the real, well-measured league-average handedness effect and treats no one as a "
     "special platoon case. Platoon is real; it is just far smaller, and far less knowable per player, "
     "than the raw split suggests.", False, False, None, None),
])
# hyperlink for THT research (appended as its own small paragraph to keep it clickable)
p_link = doc.add_paragraph()
p_link.paragraph_format.space_after = Pt(12)
add_hyperlink(p_link, "https://tht.fangraphs.com/simulating-and-identifying-platoon-players/",
              "Source: The Hardball Times, “Simulating and Identifying Platoon Players”",
              size_pt=9)

# ============================================================================
# PART III
# ============================================================================
h2("Part III", "The Mechanism: Five Per-Game Engines")

para("Time to dig into the real question: what actually moves a hitter’s output that the raw "
     "per-PA projection does not use? Each category (HR, SB, AVG, R, RBI) gets its own Poisson model, "
     "fit on 2016–2025 started games. Features fall into four groups: talent (the Steamer "
     "projection plus a heavily shrunk season-to-date component), opportunity (expected plate "
     "appearances from the batting-order slot), park (the isolated Park Factor’s HR and BACON "
     "lenses, split by the batter’s handedness), and opponent (the starting pitcher and defending "
     "team).")

para([
    ("Every table below shows the multiplier on expected production as a feature moves from its 10th "
     "to its 90th percentile, in two columns. ", False, False, None, None),
    ("Ceiling", True, False, "111111", None),
    (" is the same-season, best-case read: it scores each game as if you already knew the "
     "opponent’s full-season numbers. ", False, False, None, None),
    ("Live", True, False, "111111", None),
    (" is what the deployed tool actually delivers, using only what is knowable at pick time. The gap "
     "between the columns is the single most useful thing on the page: features whose two columns match "
     "are knowable on opening day; features whose live column collapses are the ones the season has to "
     "fill in. That collapse is not noise. It is the defense-independent-pitching insight surfacing on "
     "its own, category after category: a pitcher’s outcome rates (home runs and hits allowed) "
     "predict beautifully within a season and evaporate at any usable lag, while his skill rates "
     "(strikeouts, holding runners) carry forward.", False, False, None, None),
])

h3("Home runs: batter and park travel; the pitcher does not")
table(["Feature", "Live", "Ceiling", "Reading"],
      [
          ["In-season power update", "×1.32", "×1.32", "hot power is partly real, and stable"],
          ["Park HR factor (iPF, hand-split)", "×1.25", "×1.18", "knowable day one"],
          ["Expected PA (slot)", "×1.12", "×1.12", "more chances"],
          ["Platoon advantage", "×1.05", "×1.05", "small (league-average effect)"],
          ["Opposing SP HR rate", "×1.05", "×1.57", "huge in hindsight, dead live (DIPS)"],
          ["Opposing SP K rate", "×0.89", "×0.93", "strikeouts deny contact; a stable skill"],
      ], [2.1, 0.9, 0.9, 2.3])
caption("Scale check: within the top-150 power tier, the projection alone separates top-decile from "
        "bottom-decile realized HR rate by about 1.44×. The persistent levers (the batter’s "
        "own power trend, the park, the slot) are the same size as that talent spread and, crucially, "
        "are knowable in April. The one enormous ceiling effect (a homer-prone starter, ×1.57) is "
        "exactly the feature that cannot be known live: last year’s HR rate does not predict this "
        "year’s, so the live column falls to ×1.05. Park factors come from the isolated Park "
        "Factor leaderboard, now split by batter handedness; the split matters almost entirely for home "
        "runs (a lefty pull hitter in a short-porch park), which is why AVG uses a nearly hand-neutral "
        "BACON factor.")

h3("Stolen bases: opportunity is the constant, matchup the variable")
table(["Feature", "Live", "Ceiling", "Reading"],
      [
          ["In-season SB-rate update", "×1.65", "×1.61", "the green light, and it is stable"],
          ["Opposing SP SB-allowed rate", "×1.21", "×1.69", "strong in hindsight, mild live"],
          ["Expected PA (slot)", "×1.20", "×1.20", "more times on base, more chances"],
          ["Team SB defense", "×1.13", "×1.13", "a friendly staff overall"],
          ["Market SB prior (ADP)", "×1.09", "×1.09", "the market prices scarce speed"],
          ["Pitcher hold (runs prevented)", "×0.88", "×0.80", "a good holder deters attempts"],
          ["Left-handed pitcher", "×0.90", "×0.92", "harder to run on"],
      ], [2.1, 0.9, 0.9, 2.3])
caption("Scale check: the SB tier is the one genuinely wide talent tier — top decile steals at "
        "2.7× the bottom’s rate — because base-stealers are a distinct population, not a "
        "slice of a continuum.")
para([
    ("Era note: ", True, False, CAVEAT.__str__() if False else "B5762E", None),
    ("these effects are measured over 2016–2025, which understates them for today. The 2023 rules "
     "(bigger bases, pickoff limits) sharply raised both steal attempts and the size of the "
     "pitcher/park matchup; on the post-2023 seasons alone, the opportunity and pitcher-hold levers run "
     "materially larger than the full-window figures shown here.", False, False, None, None),
], space_after=8)
para("A separate mechanism study is worth stating: decomposing steals into attempts and successes, the "
     "pitcher governs whether a runner tries (his delivery is visible from first base) while the "
     "catcher governs whether the attempt succeeds. The live engine therefore drops the catcher (a "
     "future game’s starting catcher is unknowable at pick time) and keeps the pitcher and "
     "team-level running-game defense, which are knowable.")
p_link2 = doc.add_paragraph(); p_link2.paragraph_format.space_after = Pt(12)
r = p_link2.add_run("Sources: Baseball Savant "); r.font.size = Pt(9)
add_hyperlink(p_link2, "https://baseballsavant.mlb.com/leaderboard/pitcher-running-game",
              "pitcher running game", size_pt=9)
r2 = p_link2.add_run(" and "); r2.font.size = Pt(9)
add_hyperlink(p_link2, "https://baseballsavant.mlb.com/leaderboard/catcher-throwing",
              "catcher throwing", size_pt=9)
r3 = p_link2.add_run(" leaderboards."); r3.font.size = Pt(9)

h3("Batting average: the hardest category, where park earns its keep")
table(["Feature", "Live", "Ceiling", "Reading"],
      [
          ["Expected PA (slot)", "×1.13", "×1.13", "more chances, stable"],
          ["Park BACON factor (iPF)", "×1.06", "×1.05", "hits-on-contact environment, knowable"],
          ["In-season average update", "×1.05", "×1.05", "a whisper, correctly discounted"],
          ["Opposing SP hit rate", "×1.03", "×1.29", "hittable in hindsight, mild live (DIPS)"],
          ["Opposing SP K rate", "×0.92", "×1.00", "strikeouts deny contact; a stable skill"],
      ], [2.1, 0.9, 0.9, 2.3])
caption("Scale check: within the average tier the projection barely separates realized rate at all "
        "(about 1.14× top to bottom decile), which is both why the category is so frustrating and "
        "why environment (park, slot, opposing staff) rivals talent outright. Almost nothing separates "
        "weekly batting average except context, and the persistent parts of that context are knowable "
        "early — which is why the weekly AVG edge (66 percent of weeks) is, perhaps surprisingly, "
        "the largest of the three. Runs and RBI use their own models, shown next as context; they carry "
        "no independently validated streaming edge.")

h3("Runs and RBI: lineup role, priced honestly")
table(["Feature", "R live / ceiling", "RBI live / ceiling", "Reading"],
      [
          ["Leadoff slot (1–2)", "×1.40 / 1.42", "×1.18 / 1.20", "top of the order scores"],
          ["Middle slot (3–5)", "×1.19 / 1.20", "×1.30 / 1.33", "the heart drives in"],
          ["Own team offense", "×1.05 / 1.23", "×1.03 / 1.22", "strong live degradation (to-date)"],
          ["Opposing runs allowed", "×1.17 / 1.32", "×1.16 / 1.32", "a soft opponent, but fills in over the year"],
      ], [1.7, 1.5, 1.5, 1.9])
caption("Runs and RBI are driven by lineup role (a leadoff hitter scores, a cleanup hitter drives in), "
        "which is stable and knowable, plus the run environment of the game, which follows the same "
        "DIPS pattern: the same-season team rates are strong (×1.2–1.3) and their live "
        "versions phase in from a thinner start. These models exist to make the Overall blend coherent, "
        "not to be streamed on their own.")

# ============================================================================
# PART IV
# ============================================================================
h2("Part IV", "The Granularity Question: What Survives a Week?")

para("The single most consequential finding in this project concerns aggregation. A matchup is a "
     "property of one game; a week mixes six of them. Whether an edge survives the mixing depends "
     "entirely on whether its driver persists across the week or changes nightly.")

para("Stolen bases are the clearest case. Sorting runner-games by matchup quality (after controlling "
     "for talent), the best decile steals at 2.6 times the rate of the worst: one of the strongest "
     "single-game signals in the project (the matchup block clears a likelihood-ratio test at "
     "p < 10⁻¹²²). But the matchup component dilutes as games aggregate: 2.6× at "
     "a single game, 2.2× at a series, 2.1× at a half-week, 1.9× across a full week. Green "
     "lights and red lights average toward the middle as the week fills in, and early versions of the "
     "tool that leaned on matchup alone saw their weekly steal edge wash out entirely.")

para([
    ("The shipped engine survives the week anyway, for a reason worth stating precisely: ",
     False, False, None, None),
    ("the weekly stolen-base edge is carried by opportunity and the green light, not by the nightly "
     "matchup.", True, False, "111111", None),
    (" A leadoff hitter who has started running gets more plate appearances and more attempts every "
     "day of the week; those do not dilute, and a per-PA projection structurally cannot see either. "
     "The full picture across categories:", False, False, None, None),
])

table(["Top-5 win rate vs projection", "Game", "Half-week", "Week"],
      [
          ["Stolen bases", "61%", "66%", "62%"],
          ["Home runs", "60%", "59%", "61%"],
          ["Batting average", "54%", "62%", "66%"],
      ], [2.6, 1.3, 1.3, 1.3])
caption("Share of weeks the engine’s top five out-produced the projection’s top five, "
        "2016–2025. The categories move in opposite directions with horizon, each for a legible "
        "reason. Stolen bases peak at short horizons (matchup-rich) and hold at the week through "
        "opportunity. Batting average is the mirror image: a single game of AVG is close to pure noise "
        "(54 percent, barely above a tie), but its drivers (park, slot, opposing staff) are persistent, "
        "so the edge compounds as games accumulate. Home runs sit in between: the biggest single lever "
        "in hindsight (the opposing starter) is dead live, while park and the batter’s own power "
        "persist, leaving the edge roughly flat across horizons.")

# ============================================================================
# PART V
# ============================================================================
h2("Part V", "The Live Framework: April Floor, Summer Ceiling")

para("The two columns in Part III are the endpoints of a phase-in. In April, the live column is "
     "thinner still: with almost no current-season data, the tool leans entirely on the stable features "
     "(park, slot, handedness, the projection) and the opponent’s prior-year skills. As the season "
     "accumulates, the to-date rates fill in and the live column climbs toward its full-season value. "
     "It never quite reaches the ceiling, because the biggest ceiling levers (a pitcher’s HR and "
     "hit rates) are the ones that do not carry year to year at all.")

para("The consequence is a real phase-in, reported rather than hidden. For home runs, daily selection "
     "lift over the projection runs from a not-yet-significant +6.8 percent in April–May to a "
     "solid +15.3 percent in August–September (2022–2025 sample). Batting average is the "
     "exception that proves the rule: its edges (park, slot, opposing K rate) are knowable on day one, "
     "so its lift is roughly flat across the season. The tool earns its keep for home runs and steals "
     "as the season gives it data; for average it is ready in April.")

# ============================================================================
# PART VI
# ============================================================================
h2("Part VI", "Does It Beat the Alternatives, Where It Counts?")

para("The decision-relevant test is not abstract accuracy; it is whether the model’s picks "
     "out-produce the picks made another way. Here is the full ladder, sober number first, for each "
     "category’s pool ranked by the engine versus the projection.")

table(["Weekly", "Single pick (gap-10)", "Top-5 win rate", "Production lift"],
      [
          ["Stolen bases", "53% (proj 52.6)", "62% [57, 67]", "+13.0% [+8, +18]"],
          ["Home runs", "53% (proj 52.8)", "61% [56, 66]", "+10.8% [+7, +15]"],
          ["Batting average", "52% (proj 50.6)", "66% [61, 72]", "+11 pts [+5, +17]"],
      ], [1.8, 1.7, 1.6, 1.6])
caption("Three views of the same edge, and each is honest about a different thing. The single-pick "
        "column is the sobering truth: between two hitters ten ranks apart, the model is right about "
        "52 to 53 percent of the time, only a hair above the projection alone (shown in parentheses). "
        "The win-rate column is what happens when you stack five of those picks: the small per-pick "
        "edge concentrates on the contested slots and tips the bundle in 61 to 66 percent of weeks. The "
        "production column is the payoff in the currency you actually bank: roughly 11 to 13 percent "
        "more steals or homers, or 11 more points of average, on your streamed slots.")

h3("Why the single-pick and win-rate columns disagree (and both are right)")
para("The projection sits inside the engine (it is the prior), so for most pairs of hitters the two "
     "rankings agree and contribute identically to both scores; a pairwise-accuracy average dilutes the "
     "edge across a sea of agreement, which is why it hugs 52 percent. The win rate conditions on "
     "disagreement: when the engine’s five and the projection’s five differ by a name, the "
     "whole comparison rides on that contested slot, which is where the production lift lives. It is "
     "card-counting arithmetic: a one-percent edge per hand is invisible per hand and decisive over a "
     "season. There is also a hard ceiling; single-week outcomes are random enough that even an oracle "
     "could not push pairwise accuracy far past the mid-50s, so small movements in that column are "
     "large in relative terms. The edge is robust to the cutoff, incidentally: sweeping the selection "
     "from the top 1 to the top 10 picks moves the win rates only a few points and never reverses them.")

h3("The streamable pool, and a note on the market")
para("Everything above pools all startable hitters, including locked-in stars no one streams. "
     "Restricting both strategies to the players actually available in a typical league (outside the "
     "top 150 picks, or undrafted) answers the question as lived: stolen bases hold at 65 percent of "
     "weeks [59, 70] and home runs at 63 percent [58, 68], while batting average drops to 53 percent "
     "[47, 60] and does not clear significance. The average edge lives among rostered players, where "
     "park and lineup context separate real talent; on the waiver tier, where the talent spread in "
     "average is nearly flat, a week of AVG is noise most of the way down. Treat the AVG column as a "
     "tiebreaker among rosterable bats, not a deep-pool discovery engine. As for the market: preseason "
     "ADP, given the projection, is a mild negative — draft-day enthusiasm is a slight fade once "
     "the projection is known — with one exception, steals, where the market prices speed the "
     "projection under-weights. ADP therefore earns a small, real place in the SB engine and nowhere "
     "else.")

# ============================================================================
# PART VII  -- Pace Standard
# ============================================================================
h2("Part VII", "The Pace Standard: What Counts as a Good Week")

para("Beating the projection is a relative claim; a streamer also needs an absolute one: is this slot "
     "keeping the pace that wins the category? The pace is derivable from real standings data. The "
     "80th percentile of NFBC Main Event season standings requires roughly 307 HR, 183 SB, 1,058 R, "
     "1,032 RBI, and a .2564 average. A season contains 27.6 Mon–Sun scoring weeks (measured from "
     "the schedule, 2016–2025) and a Main Event lineup starts 14 hitters, so the winning pace per "
     "lineup slot per week is about 0.79 HR, 0.47 SB, 2.7 R, and 2.7 RBI. Weeks come in integers, which "
     "sets the natural thresholds: a good week is one home run or more, one steal or more (two or more "
     "is the green-light week), three runs, three RBI.")

para([
    ("One caveat is required and worth keeping: ", True, False, "B5762E", None),
    ("the pace is a team-sum requirement, and your stars bank surplus. The honest reading of a slot "
     "threshold is “this slot kept the winning pace this week,” not “every hitter must "
     "clear it every week.” Batting average is deliberately excluded from the binary treatment: a "
     "weekly average bar mechanically favors fewer at-bats (variance clears bars that talent cannot), "
     "which is backwards for a streaming tool, so average is scored as pace surplus — hits banked "
     "above a .2564 pace on the at-bats taken.", False, False, None, None),
])

_p_formula = doc.add_paragraph()
_p_formula.paragraph_format.space_after = Pt(10)
_r1 = _p_formula.add_run(
    "Because the engines are per-game count models, every good-week probability is a direct "
    "arithmetic consequence of predictions the tool already makes: P(at least one home run) = 1 − e")
_r_sup = _p_formula.add_run("−λ")
_r_sup.font.superscript = True
_r2 = _p_formula.add_run(
    ", with λ the week’s summed per-game expectation. Ranking by that "
    "probability and ranking by expected production select identical players, so this section changes "
    "how the tool talks, not what it picks. The question is whether those probabilities can be taken "
    "literally, and that is a testable claim: across out-of-sample decile bins, predicted versus "
    "realized good-week frequency tracks with intercepts near zero and slopes of 1.05 to 1.08 for all "
    "four thresholds (HR and SB, one-plus and two-plus). In words: when the engine says 60 percent, it "
    "happens about 62 percent of the time — honestly calibrated, leaning slightly conservative. ")
para("This doubles as the direct test of the per-game Poisson architecture (correlated games within a "
     "week would surface here as overconfidence; the opposite appeared), and the probabilities beat "
     "projection-only probabilities on Brier score at all four thresholds, by small but decisively real "
     "margins.")

table(["Streamable pool, weekly", "Engine picks: good weeks", "Projection picks: good weeks",
       "Difference", "Pace surplus (engine vs proj)"],
      [
          ["HR week (1+)", "76%", "71%", "+4.3 [+2.2, +6.3]", "+0.69 vs +0.52 HR"],
          ["SB week (1+)", "68%", "61%", "+6.8 [+4.4, +9.2]", "+0.83 vs +0.64 SB"],
          ["Green-light week (2+ SB)", "35%", "29%", "+6.2 [+4.0, +8.2]", "—"],
          [("AVG (pace surplus only)", 2), "+1.00 vs +0.71 hits above pace", "+0.29 [+0.18, +0.41]", "—"],
      ], [1.9, 1.4, 1.4, 1.5, 1.6])
caption("Share of top-five picks (players outside the top 150 draft picks, or undrafted) that "
        "delivered a good week, 2016–2025 out of sample, week-clustered intervals; pace surplus is "
        "realized production above the 80th-percentile pace per streamed slot-week. Both "
        "strategies’ picks beat pace comfortably (the dynamic pool surfaces in-season risers); the "
        "engine’s picks beat it more often and by more in every category. A nuance worth noticing: "
        "on this pool the weekly AVG win rate did not clear significance (Part VI), but the AVG pace "
        "surplus does — the continuous metric keeps the magnitude information a binary comparison "
        "discards. Both statements are true; they answer different questions.")

para("Why the pace standard stays a reporting lens rather than the primary metric: roto is cumulative, "
     "and every homer past the first still counts. A three-homer week banks three, and a binary label "
     "would score it identically to a one-homer week, which is precisely the information the "
     "production-based validation preserves. The thresholds here are the Main Event’s; a 12-team "
     "format lowers the bar and an OBP format changes the average column entirely, so the paces travel "
     "with the format, not with the tool.")

# ============================================================================
# PART VIII
# ============================================================================
h2("Part VIII", "The Overall Blend: Reliability-Shrunk Stacking")

para("Summing five category scores into one is harder than it sounds. Only three categories carry a "
     "validated edge, yet all five carry a full category of weekly noise, so an equal-weight blend is "
     "structurally diluted; in early testing it did not clear the projection with confidence (a plain "
     "blend wins 56 percent of weeks). That is expected arithmetic, not a bug: signal from three "
     "places, noise from five.")

para("The fix is reliability-shrunk stacking. Each category’s prediction is split into its "
     "projection baseline plus the model’s deviation, and a coefficient learned per category (on "
     "training seasons only) decides how much of that deviation to trust before summing. The weights "
     "came out remarkably stable across nine held-out seasons:")

table(["Category", "Trust in model deviation", "Read"],
      [
          ["Home runs", "0.53", "trust it"],
          ["Runs", "0.53", "context is reliable"],
          ["RBI", "0.45", "mostly"],
          ["Stolen bases", "0.24", "discount"],
          ["Batting average", "0.21", "discount hardest"],
      ], [2.0, 2.2, 2.3])
caption("Single-week AVG and SB deviations are the noisiest, so they are shrunk hardest; home runs are "
        "trusted most. The quiet surprise is runs, trusted as much as homers: not because runs are "
        "independently streamable, but because they are predictable from batting-order slot and team "
        "offense, which are stable. Trustworthy and edge-bearing are different properties, and the "
        "stack respects the difference — which is exactly why the Part III R and RBI tables show "
        "large, stable lineup-role effects.")

para("Fit out of sample over nine seasons (91,097 hitter-weeks), the shrunk Overall beats the "
     "projection’s top five in 61 percent of weeks [55, 67], up from 56 percent for the plain "
     "blend and clearly clear of a coin flip. The category-weight sliders in the app sit on top of "
     "this: punt average and crank steals, and the blend re-weights to the standings you are actually "
     "chasing, which is the one form of value no generic backtest can measure.")

# ============================================================================
# PART IX
# ============================================================================
h2("Part IX", "Honest Limitations")

box_p = doc.add_paragraph()
box_p.paragraph_format.space_after = Pt(4)
r = box_p.add_run("WHAT THE TOOL DOES NOT YET DO")
r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x8a, 0x50, 0x15)

caveat_item("Playing-time risk.", "The engine conditions on “he starts.” A bench and "
            "platoon risk layer (will he start?) is a planned addition, not a current feature.")
caveat_item("The steal numbers are conservative.", "All SB effects here are measured over "
            "2016–2025. The 2023 rule changes strengthened the steal matchup materially, so the "
            "current-era pitcher-hold and opportunity effects run larger than the full-window figures "
            "shown; the tool is, if anything, understated for steals today.")
caveat_item("Late-week probables.", "Starters are firm roughly five days out and reasonably stable to "
            "ten; a probable not yet listed is scored with a neutral matchup and flagged in the "
            "schedule cell.")
caveat_item("Catcher identity.", "A future game’s starting catcher is genuinely unknowable at pick "
            "time, so the live engine drops the catcher feature; the team-level running-game defense, "
            "which is knowable, captures part of it. Conveniently, the catcher acts on execution rather "
            "than deterrence, so the omission costs less than it appears.")
caveat_item("Weather.", "Temperature moves home runs materially and is not yet a feature. There is "
            "deliberately no SB park factor: the bases are 90 feet apart everywhere.")
caveat_item("Waiver-tier batting average.", "As shown in Part VI, the AVG edge does not clear "
            "significance on the deep-availability pool. Treat the AVG column as a tiebreaker among "
            "established bats, not a discovery engine.")
doc.add_paragraph().paragraph_format.space_after = Pt(4)

para("Two methodological notes for the skeptical reader. Park factors are estimated across the full "
     "era, which leaks a small amount of future information into the backtest; the effect is bounded "
     "and the live tool uses only published factors. And the many out-of-sample tests here were not "
     "pre-registered; the findings that survive do so with mechanistic explanations and cross-category "
     "replication (the DIPS pattern appears independently in HR, AVG, R, and RBI), and borderline "
     "results are treated as noise rather than near-misses throughout.")

# ============================================================================
# PART X
# ============================================================================
h2("Part X", "Conclusions: When and How to Use It")

para([
    ("Use the category columns for the category you need. ", True, False, "111111", None),
    ("The validated edges are category-specific, and so is streaming: the SB column for a steals "
     "deficit, the HR column for power, the AVG column as a tiebreaker among rosterable bats (not for "
     "deep-pool darts). The Overall column, with the weight sliders set to match league standings, is "
     "the default sort and carries its own validated edge (61 percent of weeks).", False, False, None, None),
])

para([
    ("Match the horizon to the category. ", True, False, "111111", None),
    ("Stolen-base matchups are a daily and half-week weapon; the weekly SB read leans on opportunity "
     "and the green light. Batting average is the opposite: nearly useless for a single day (54 "
     "percent), strongest across a full week (66). Home runs work at any horizon. Setting a Monday "
     "lineup for the week and picking a Friday spot-add are both supported, and they lean on different "
     "parts of the engine.", False, False, None, None),
])

para([
    ("Trust it more in August than in April — except for average. ", True, False, "111111", None),
    ("The HR and SB engines phase in as to-date rates accumulate; early-season output leans on stable "
     "priors and is honest about its thinner edge. The AVG engine, built on park and slot and "
     "opposing-staff skill, is ready on opening day.", False, False, None, None),
])

para([
    ("Read the probabilities literally, and stream against the pace. ", True, False, "111111", None),
    ("The good-week probabilities are calibrated out of sample (a stated 60 percent happens about 62 "
     "percent of the time), and the streamed slot’s job is concrete: keep the 80th-percentile "
     "winning pace of roughly 0.8 HR and 0.5 SB per slot-week. The engine’s waiver-tier picks kept "
     "that pace more often than the projection’s in every category tested, and by more.",
     False, False, None, None),
])

para([
    ("The headline findings, in one place. ", True, False, "111111", None),
    ("Projections explain under nine percent of week-to-week hitter value; among similar hitters they "
     "pick the better week at roughly a coin flip, and the models add only a hair to that per decision. "
     "That hair compounds: matchup-aware per-game modeling beats the projection in three of five weeks "
     "per category and produces measurably more of each stat on the slots it changes. The edges have "
     "legible mechanics: opportunity and the green light carry steals, persistence carries average, the "
     "batter and the park carry home runs, and the opposing pitcher’s outcome rates carry almost "
     "nothing that survives to next season. Recency ranking, the most popular streaming heuristic, is "
     "the worst strategy tested. Platoon effects are real at the league level and small once the "
     "projection is known. The market’s residual opinion is a slight fade, except about steals.",
     False, False, None, None),
])

# ---- footer -----------------------------------------------------------
fp = doc.add_paragraph()
fp.paragraph_format.space_before = Pt(18)
pPr = fp._p.get_or_add_pPr()
pbdr = OxmlElement("w:pBdr")
top = OxmlElement("w:top")
top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "6"); top.set(qn("w:color"), "E8E4D8")
pbdr.append(top); pPr.append(pbdr)
r = fp.add_run(
    "CollinMyShot FBB Tools · Hitter Streamonator · Methodology · Validation: "
    "leave-one-season-out, 2016–2025 (2020 excluded), ~368,000 hitter-games; selection win rates "
    "and production lifts clustered by week. Effect sizes and win rates share one provenance (shipped "
    "per-game specs, hand-split iPF park). Pace standard: NFBC Main Event 80th-percentile season "
    "standings targets, divided across measured Mon–Sun scoring weeks and 14 starting slots. "
    "Platoon grids: FanGraphs splits, 2015–2025. Platoon stabilization: The Hardball Times. "
    "Running-game metrics: Baseball Savant.")
r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GRAY

out_path = "/private/tmp/claude-501/-Users-ckaufman-Documents-New-project/57b453dd-83f5-4aa1-abf9-e40500bf9c39/scratchpad/Hitter_Streamonator_Methodology.docx"
doc.save(out_path)
print("saved:", out_path)
